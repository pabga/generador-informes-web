import streamlit as st
import pandas as pd
from docx import Document
import io
from datetime import datetime
import gspread


# --- FUNCIÓN PARA CREAR UNA TABLA (REUTILIZABLE) ---
def crear_tabla_en_documento(documento, marcador_tabla, dataframe_personas):
    # (Esta función no cambia)
    parrafo_marcador = None
    for p in documento.paragraphs:
        if marcador_tabla in p.text:
            parrafo_marcador = p
            break

    if parrafo_marcador:
        parrafo_marcador.text = ""
        tabla = documento.add_table(rows=1, cols=3)
        try:
            tabla.style = 'Tabla con cuadrícula'
        except KeyError:
            try:
                tabla.style = 'Table Grid'
            except KeyError:
                st.warning(f"Estilo de tabla no encontrado para {marcador_tabla}.")
                pass

        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Jerarquía'
        hdr_cells[1].text = 'DNI'
        hdr_cells[2].text = 'Nombre y Apellido'

        for index, persona in dataframe_personas.iterrows():
            row_cells = tabla.add_row().cells
            row_cells[0].text = str(persona.get('Jerarquia', ''))
            row_cells[1].text = str(persona.get('DNI', ''))
            row_cells[2].text = str(persona.get('Nombre_Apellido', ''))

        parrafo_marcador._p.addnext(tabla._element)
    else:
        # Ya no mostramos advertencia si la plantilla no usa la tabla
        pass


# --- CONEXIÓN A GOOGLE SHEETS (NO CAMBIA) ---
try:
    gc = gspread.service_account_from_dict(st.secrets["gcp_service_account"])
    sh_personas = gc.open("Base de Datos - Personas").sheet1
    sh_docentes = gc.open("Base de Datos - Docentes").sheet1
    sh_cursos = gc.open("Base de Datos - Cursos").sheet1
    df_personas_full = pd.DataFrame(sh_personas.get_all_records())
    df_docentes_full = pd.DataFrame(sh_docentes.get_all_records())
    df_cursos = pd.DataFrame(sh_cursos.get_all_records())
    df_personas_full['DNI'] = df_personas_full['DNI'].astype(str)
    df_docentes_full['DNI'] = df_docentes_full['DNI'].astype(str)
except Exception as e:
    st.error(f"Error al conectar con Google Sheets: {e}")
    st.stop()


# --- FUNCIÓN PRINCIPAL (AHORA RECIBE EL TEXTO MANUAL) ---
def generar_documento(curso_elegido_df, dnis_participantes, plantilla_a_usar, texto_manual):  # <-- PARÁMETRO NUEVO
    participantes_df = df_personas_full[df_personas_full['DNI'].isin(dnis_participantes)]
    st.info(f"Se encontraron {len(participantes_df)} de {len(dnis_participantes)} participantes.")

    dnis_docentes_str = str(curso_elegido_df.get('DNI_Docentes', ''))
    dnis_docentes = [dni.strip() for dni in dnis_docentes_str.split(',') if dni.strip()]
    docentes_df = df_docentes_full[df_docentes_full['DNI'].isin(dnis_docentes)]
    st.info(f"Se encontraron {len(docentes_df)} docentes para este curso.")

    documento = Document(plantilla_a_usar)

    # Combinamos los datos del curso con el texto manual
    datos_completos = curso_elegido_df.to_dict()
    datos_completos['S_D'] = texto_manual  # <-- AÑADIMOS EL TEXTO MANUAL AL DICCIONARIO

    todos_los_parrafos = list(documento.paragraphs)
    for tabla in documento.tables:
        for row in tabla.rows:
            for cell in row.cells:
                for parrafo in cell.paragraphs:
                    todos_los_parrafos.append(parrafo)

    for parrafo in todos_los_parrafos:
        for key, value in datos_completos.items():
            marcador = f"{{{{{key}}}}}"
            if marcador in parrafo.text:
                dato_a_reemplazar = ''
                if key in ['Fecha_Inicio', 'Fecha_Fin'] and pd.notna(value) and value != '':
                    try:
                        dato_a_reemplazar = pd.to_datetime(value).strftime('%d/%m/%Y')
                    except (ValueError, TypeError):
                        dato_a_reemplazar = str(value)
                else:
                    dato_a_reemplazar = str(value)
                parrafo.text = parrafo.text.replace(marcador, dato_a_reemplazar)

    crear_tabla_en_documento(documento, '{{TABLA_PARTICIPANTES}}', participantes_df)
    crear_tabla_en_documento(documento, '{{TABLA_DOCENTES}}', docentes_df)

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer


# --- INTERFAZ DE LA PÁGINA WEB (CON CAMPO DE TEXTO) ---
st.title("Generador de Informes de Cursos 🚀")

plantillas_disponibles = {
    "Memo Autogestionado": "memo_autogestionado.docx",
    "Memo Autogestionado (con clase)": "memo_autogestionado_con_clase.docx",
    "Memo Extension" : "memo_extension.docx"
}
opcion_plantilla_display = st.selectbox(
    "Paso 1: Seleccione el tipo de documento",
    list(plantillas_disponibles.keys())
)
plantilla_seleccionada_archivo = plantillas_disponibles[opcion_plantilla_display]

lista_cursos = df_cursos['Nombre_Curso'].tolist()
curso_seleccionado_nombre = st.selectbox("Paso 2: Seleccione el curso", lista_cursos)
archivo_dni_subido = st.file_uploader("Paso 3: Suba el archivo `lista_dni.txt` con los DNI de los participantes",
                                      type="txt")

# --- NUEVO: CAMPO DE TEXTO PARA CONTENIDO MANUAL ---
st.markdown("---")  # Una línea para separar
texto_adicional = st.text_area(
    "Paso 4 (Opcional): Agregue texto ej 000000/AGO/25",
    height=150  # Altura del cuadro de texto
)

if st.button("Generar Documento"):
    if archivo_dni_subido is not None:
        with st.spinner('Procesando...'):
            dnis = archivo_dni_subido.getvalue().decode("utf-8").splitlines()
            dnis_limpios = [dni.strip() for dni in dnis if dni.strip()]
            curso_elegido_df = df_cursos[df_cursos['Nombre_Curso'] == curso_seleccionado_nombre].iloc[0]

            # Pasamos el texto adicional a la función
            buffer_documento = generar_documento(curso_elegido_df, dnis_limpios, plantilla_seleccionada_archivo,
                                                 texto_adicional)

            nombre_curso_corto = curso_seleccionado_nombre.replace(" ", "_")[:20]
            nombre_archivo = f"{opcion_plantilla_display.split(' ')[0]}_{nombre_curso_corto}_{datetime.now().strftime('%Y-%m-%d')}.docx"

            st.success("¡Documento generado con éxito!")
            st.download_button(
                label="Descargar Documento Word",
                data=buffer_documento,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Por favor, suba el archivo `lista_dni.txt`.")